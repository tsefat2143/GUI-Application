import docx
import os

os.getcwd()

doc = docx.Document('template.docx')
print(doc.paragraphs[0].text)
addPar = doc.add_paragraph('Add text')
doc.save('template.docx')
